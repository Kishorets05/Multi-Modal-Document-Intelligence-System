from pathlib import Path
from typing import Any

from docx import Document


class DocxReader:
    """Extract metadata from a DOCX document."""

    def __init__(self, file_path: Path) -> None:
        self.file_path = Path(file_path)

    def read(self) -> dict[str, Any]:
        """Return extracted DOCX metadata."""
        document = Document(self.file_path)
        paragraphs = [p for p in document.paragraphs if p.text.strip()]
        heading_count = sum(
            1
            for p in document.paragraphs
            if p.style.name.lower().startswith("heading")
        )

        return {
            "document_type": "docx",
            "page_count": None,
            "paragraph_count": len(paragraphs),
            "heading_count": heading_count,
            "image_width": None,
            "image_height": None,
            "document_size": self.file_path.stat().st_size,
            "paragraphs": [p.text.strip() for p in paragraphs],
            "document_metadata": {
                "core_properties": {
                    name: getattr(document.core_properties, name)
                    for name in (
                        "author",
                        "category",
                        "comments",
                        "created",
                        "identifier",
                        "keywords",
                        "language",
                        "last_modified_by",
                        "last_printed",
                        "modified",
                        "revision",
                        "subject",
                        "title",
                    )
                    if hasattr(document.core_properties, name)
                }
            },
            "reader_used": self.__class__.__name__,
        }
